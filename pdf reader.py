from docx.api import Document
import docx.api
from  docx.text.paragraph import Paragraph
print(dir(docx))
from docx.table import _Cell, Table
# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
document = Document('C:/Users/eyyup/Desktop/kelimeler.docx')
print(dir(document))
print(dir(docx.api))
tables = document.tables
ph = document.paragraphs[0]
print(dir(tables))
print(tables)
print(len(tables))
data = []
b = tables[1]
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """

    if isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        yield Paragraph(child, parent)

for b in tables[30:]:
    try:
        for row in b.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    data.append(paragraph.text)
    except Exception as e:
        for tc in b._tbl.iter_tcs():
            cell = _Cell(tc, b)
            for b_tc in iter_block_items(cell):
                if isinstance(b_tc, Paragraph):
                    data.append(b_tc.text)
#print(data)
jump_counter=0
sayac=1
index=0
words =[]
for index in range(len(data)):

    if data[index]==str(sayac) or data[index]==str(sayac+1) or data[index]==str(sayac+2) or data[index]==str(sayac+3):
        if  data[index]!=str(sayac):
             jump_counter+=int(data[index])-sayac
        if (sayac>406 and sayac<415) or sayac==782 or sayac==842:
            print(data[index:index+20])
        adim=1
        while True:
            if data[index+adim].lower()!=str(sayac) and data[index+adim].lower()!='':
                if data[index+adim].lower()=="378":
                    print("geldi:",sayac)
                words.append(data[index+adim].lower())
                break
            else:
                adim+=1
        
        sayac=int(data[index])+1
    else:
        try:
            int_deger=int(data[index])
            if (int_deger>406 and int_deger<415) or int_deger==782 or int_deger==842:
                print(data[index:index+20])
            if int_deger>sayac:
                
                adim=1
                while True:
                    if data[index+adim].lower()!=str(int_deger) and data[index+adim].lower()!='':
                        words.append(data[index+adim].lower())
                        break
                    else:
                        adim+=1
                sayac=int_deger+1
        except:
            None
print("words:",words)
print("jump:",jump_counter)
print(len(words))
dosya=open("C:/Users/eyyup/Desktop/sözlük/eklenecekler.txt","w")
for i in words:
    try:
        int(i)
        continue
    except:
        None
    if "(" not in i and ")" not in i:
        dosya.write(i+"\n")
dosya.close()
